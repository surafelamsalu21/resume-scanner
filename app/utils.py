import re
import logging
from typing import Dict, List, Optional, Tuple
from datetime import datetime
import json
import hashlib
from openai import OpenAI
from os import getenv
from flask import current_app

# Try to import spacy, but don't fail if it's not available
try:
    import spacy
    from spacy.language import Language
    SPACY_AVAILABLE = True
except ImportError:
    SPACY_AVAILABLE = False
    from typing import Any
    Language = Any  # Placeholder type

# Configure logging
logger = logging.getLogger(__name__)


class ResumePreprocessor:
    """Handles resume text preprocessing and cleaning."""

    def __init__(self):
        """Initialize the preprocessor with OpenAI client."""
        self.client = OpenAI(api_key=getenv('OPENAI_API_KEY'))
        if not getenv('OPENAI_API_KEY'):
            logger.warning("OpenAI API key not found in environment variables")

    def clean_text(self, text: str) -> str:
        """Clean and normalize resume text using OpenAI."""
        try:
            response = self.client.chat.completions.create(
                model="gpt-4o-mini",
                messages=[
                    {"role": "system", "content": """You are a specialized text preprocessing expert for resumes. Your tasks include:
                    1. Remove irrelevant special characters while preserving important ones (.,@-())
                    2. Standardize formatting (spacing, line breaks, bullet points)
                    3. Correct obvious spelling and formatting errors
                    4. Maintain professional terminology and abbreviations
                    5. Preserve contact information format
                    6. Keep date formats consistent
                    7. Standardize section headers
                    8. Remove any non-printable characters or hidden formatting
                    9. Ensure proper spacing around punctuation
                    10. Maintain list and bullet point structures"""},
                    {"role": "user", "content": text}
                ],
                temperature=0.3
            )
            return response.choices[0].message.content.strip()
        except Exception as e:
            logger.error(f"Error in text cleaning with OpenAI: {str(e)}")
            logger.info("Falling back to basic text cleaning")
            # Fallback to basic cleaning if API call fails
            text = re.sub(r'[^\w\s.,@()-]', ' ', text)
            text = re.sub(r'\s+', ' ', text)
            text = re.sub(r'\.{2,}', '.', text)
            return text.strip()

    def extract_sections(self, text: str) -> Dict[str, str]:
        """Extract common resume sections using OpenAI."""
        try:
            prompt = """You are a professional resume parser with expertise in identifying and extracting standard resume sections. Your task is to:

            1. Identify and extract the following sections with their complete content:
               - Education: Include degrees, institutions, dates, GPA, honors, relevant coursework
               - Experience: Include job titles, companies, dates, responsibilities, achievements
               - Skills: Both technical and soft skills, grouped by category
               - Projects: Personal, academic, or professional projects with descriptions
               - Certifications: Professional certifications with dates and issuing organizations

            2. For each section:
               - Maintain the original formatting and structure
               - Keep all relevant bullet points and sub-sections
               - Preserve dates and durations
               - Keep specific metrics and achievements
               - Maintain technical terminology

            3. Handle special cases:
               - Combine similar sections (e.g., "Technical Skills" and "Skills")
               - Handle non-standard section names
               - Process both chronological and functional resume formats
               - Maintain hierarchical structure within sections

            Please format the output as a JSON object with these exact keys: education, experience, skills, projects, certifications.
            If a section is empty, return an empty string for that key.

            Resume text to process:
            """

            response = self.client.chat.completions.create(
                model="gpt-4o-mini",
                messages=[
                    {"role": "system", "content": "You are an expert resume section parser with deep understanding of resume structures and formats. Always respond with valid JSON only."},
                    {"role": "user", "content": prompt + text}
                ],
                temperature=0.3
            )

            try:
                sections = json.loads(response.choices[0].message.content)
                required_sections = {'education', 'experience',
                                     'skills', 'projects', 'certifications'}
                for section in required_sections:
                    if section not in sections:
                        sections[section] = ''
                return sections
            except json.JSONDecodeError:
                logger.error("Failed to parse OpenAI response as JSON")
                return {
                    'education': '',
                    'experience': '',
                    'skills': '',
                    'projects': '',
                    'certifications': ''
                }

        except Exception as e:
            logger.error(f"Error in section extraction: {str(e)}")
            # Fallback to regex-based extraction if API call fails
            sections = {
                'education': '',
                'experience': '',
                'skills': '',
                'projects': '',
                'certifications': ''
            }

            section_patterns = {
                'education': r'(?i)(education|academic background|qualifications)',
                'experience': r'(?i)(experience|work history|employment|professional background)',
                'skills': r'(?i)(skills|technical skills|competencies)',
                'projects': r'(?i)(projects|personal projects|professional projects)',
                'certifications': r'(?i)(certifications|certificates|professional development)'
            }

            for section, pattern in section_patterns.items():
                matches = list(re.finditer(pattern, text))
                if matches:
                    start = matches[0].end()
                    next_sections = [m.start() for m in re.finditer(
                        r'(?i)(' + '|'.join(section_patterns.values()) + ')', text[start:])]
                    end = start + \
                        (next_sections[0]
                         if next_sections else len(text[start:]))
                    sections[section] = text[start:end].strip()

            return sections


class Anonymizer:
    """Handles personal information anonymization in resumes."""

    def __init__(self, nlp: Optional[Language] = None):
        self.nlp = nlp
        self.pii_patterns = {
            'email': r'\b[\w\.-]+@[\w\.-]+\.\w+\b',
            'phone': r'\b(\+\d{1,2}\s?)?\(?\d{3}\)?[\s.-]?\d{3}[\s.-]?\d{4}\b',
            'ssn': r'\b\d{3}-\d{2}-\d{4}\b',
            'url': r'https?://(?:[-\w.]|(?:%[\da-fA-F]{2}))+[^\s]*'
        }

    def anonymize_text(self, text: str) -> Tuple[str, Dict[str, str]]:
        """
        Anonymize personal information in text while preserving structure.
        Returns anonymized text and mapping of original to anonymized values.
        """
        mapping = {}
        
        # Only use spacy if it's available and nlp model is provided
        if SPACY_AVAILABLE and self.nlp is not None:
            doc = self.nlp(text)
            
            # Anonymize named entities
            for ent in doc.ents:
                if ent.label_ in ['PERSON', 'ORG', 'GPE', 'LOC']:
                    hash_value = self._generate_hash(ent.text)[:8]
                    replacement = f"[{ent.label_}_{hash_value}]"
                    mapping[ent.text] = replacement
                    text = text.replace(ent.text, replacement)

        # Anonymize PII using regex patterns (works without spacy)
        for pii_type, pattern in self.pii_patterns.items():
            matches = re.finditer(pattern, text)
            for match in matches:
                original = match.group()
                hash_value = self._generate_hash(original)[:8]
                replacement = f"[{pii_type.upper()}_{hash_value}]"
                mapping[original] = replacement
                text = text.replace(original, replacement)

        return text, mapping

    @staticmethod
    def _generate_hash(text: str) -> str:
        """Generate a consistent hash for text."""
        return hashlib.md5(text.encode()).hexdigest()

    def restore_text(self, anonymized_text: str, mapping: Dict[str, str]) -> str:
        """Restore original text using the anonymization mapping."""
        restored_text = anonymized_text
        reverse_mapping = {v: k for k, v in mapping.items()}
        for anon, original in reverse_mapping.items():
            restored_text = restored_text.replace(anon, original)
        return restored_text


class ResumeValidator:
    """Validates resume content and structure."""

    @staticmethod
    def validate_required_fields(sections: Dict[str, str]) -> List[str]:
        """Check for required resume sections and return missing fields."""
        required_fields = ['education', 'experience', 'skills']
        missing_fields = []

        for field in required_fields:
            if not sections.get(field):
                missing_fields.append(field)
            elif len(sections[field].split()) < 10:  # Minimum content check
                missing_fields.append(f"{field} (insufficient content)")

        return missing_fields

    @staticmethod
    def validate_date_formats(text: str) -> List[str]:
        """Validate and standardize date formats in resume."""
        date_patterns = [
            r'\b(Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]* \d{4}\b',
            r'\b\d{2}/\d{2}/\d{4}\b',
            r'\b\d{4}-\d{2}-\d{2}\b'
        ]

        invalid_dates = []
        for pattern in date_patterns:
            matches = re.finditer(pattern, text)
            for match in matches:
                try:
                    # Attempt to parse the date
                    date_str = match.group()
                    datetime.strptime(date_str, '%Y-%m-%d')
                except ValueError:
                    invalid_dates.append(date_str)

        return invalid_dates


def preprocess_resume(text: str, anonymize: bool = False) -> Dict[str, any]:
    """
    Main function to preprocess resume text.
    Returns preprocessed text and metadata.
    """
    try:
        preprocessor = ResumePreprocessor()

        # Clean text
        cleaned_text = preprocessor.clean_text(text)

        # Extract sections
        sections = preprocessor.extract_sections(cleaned_text)

        # Validate content
        validator = ResumeValidator()
        missing_fields = validator.validate_required_fields(sections)
        invalid_dates = validator.validate_date_formats(cleaned_text)

        # Anonymize if requested
        anonymized_text = None
        mapping = None
        if anonymize and SPACY_AVAILABLE:
            # Initialize spacy model if needed
            try:
                nlp = spacy.load("en_core_web_sm")
                anonymizer = Anonymizer(nlp)
                anonymized_text, mapping = anonymizer.anonymize_text(cleaned_text)
            except Exception as e:
                logger.warning(f"Could not load spacy model for anonymization: {str(e)}")
                # Fall back to regex-only anonymization
                anonymizer = Anonymizer(None)
                anonymized_text, mapping = anonymizer.anonymize_text(cleaned_text)
        elif anonymize:
            # Use regex-only anonymization if spacy is not available
            logger.warning("Spacy not available, using regex-only anonymization")
            anonymizer = Anonymizer(None)
            anonymized_text, mapping = anonymizer.anonymize_text(cleaned_text)

        return {
            'original_text': text,
            'cleaned_text': cleaned_text,
            'anonymized_text': anonymized_text,
            'anonymization_mapping': mapping,
            'sections': sections,
            'validation': {
                'missing_fields': missing_fields,
                'invalid_dates': invalid_dates,
                'is_valid': not (missing_fields or invalid_dates)
            },
            'metadata': {
                'processing_timestamp': datetime.now().isoformat(),
                'word_count': len(cleaned_text.split()),
                'has_required_sections': not bool(missing_fields)
            }
        }

    except Exception as e:
        logger.error(f"Error preprocessing resume: {str(e)}")
        raise


def validate_json_structure(data: Dict) -> bool:
    """Validate JSON structure of processed resume data."""
    required_keys = {
        'technical_assessment',
        'experience_assessment',
        'education_assessment',
        'soft_skills_assessment',
        'overall_recommendation'
    }

    try:
        # Check if all required top-level keys exist
        if not all(key in data for key in required_keys):
            return False

        # Validate specific field types
        if not isinstance(data['experience_assessment'].get('years_relevant_experience'), (int, float)):
            return False

        if not isinstance(data['education_assessment'].get('degree_relevance'), (int, float)):
            return False

        if not isinstance(data['overall_recommendation'].get('fit_score'), (int, float)):
            return False

        return True

    except (AttributeError, KeyError):
        return False


class AIProcessor:
    """Handles all AI-related processing using OpenAI models."""

    def __init__(self):
        # Initialize OpenAI API key
        self.api_key = getenv(
            'OPENAI_API_KEY') or current_app.config.get('OPENAI_API_KEY')
        self.model = getenv(
            'OPENAI_MODEL', 'gpt-4o-mini') or current_app.config.get('OPENAI_MODEL', 'gpt-4o-mini')

        # Set up logging
        self.logger = logging.getLogger(__name__)

        # Validate API key
        if not self.api_key:
            self.logger.error("OpenAI API key is not set")
            raise ValueError("OpenAI API key is not set")

        # Set up OpenAI client (new way - no need to set global api_key)
        self.client = OpenAI(api_key=self.api_key)

    def analyze_resume(self, resume_text: str) -> Dict[str, any]:
        """Analyze resume using OpenAI."""
        try:
            self.logger.info(
                f"Analyzing resume with OpenAI (model: {self.model})")
            self.logger.info(
                f"Resume text length: {len(resume_text)} characters")

            # Save resume text to file for debugging
            with open("resume_text.txt", "w") as f:
                f.write(resume_text)

            # Create prompt for OpenAI
            prompt = f"""
            Analyze the following resume and extract key information:
            
            RESUME:
            {resume_text}
            
            Please provide a detailed analysis in JSON format with the following structure:
            {{
                "candidate_name": "Full name of the candidate",
                "contact_info": {{
                    "email": "Email address",
                    "phone": "Phone number",
                    "location": "Location"
                }},
                "skills": [
                    {{"name": "Skill name", "level": "Advanced"}}
                ],
                "experience": [
                    {{
                        "title": "Job title",
                        "company": "Company name",
                        "duration": "Duration in years",
                        "description": "Brief description"
                    }}
                ],
                "education": [
                    {{
                        "degree": "Degree name",
                        "institution": "Institution name",
                        "field": "Field of study",
                        "year": "Year completed"
                    }}
                ],
                "summary": "Brief professional summary"
            }}
            
            Ensure all fields are filled based on the resume content. If information is not available, use null or empty arrays.
            """

            # Log that we're about to call OpenAI
            self.logger.info("Sending request to OpenAI API")

            # Call OpenAI API
            response = self.client.chat.completions.create(
                model=self.model,
                messages=[
                    {"role": "system", "content": "You are an expert resume analyzer. Extract accurate information from resumes and return only valid JSON."},
                    {"role": "user", "content": prompt}
                ],
                temperature=0.3,
                max_tokens=2000
            )

            # Log successful response
            self.logger.info("Received response from OpenAI API")

            # Extract and parse the response
            analysis_text = response.choices[0].message.content.strip()

            # Save the raw response to a file for debugging
            with open("openai_response.txt", "w") as f:
                f.write(analysis_text)

            # Try to extract JSON from the response
            try:
                # Find JSON in the response if it's wrapped in markdown or other text
                json_start = analysis_text.find('{')
                json_end = analysis_text.rfind('}') + 1
                if json_start >= 0 and json_end > json_start:
                    json_str = analysis_text[json_start:json_end]
                    analysis_data = json.loads(json_str)
                else:
                    analysis_data = json.loads(analysis_text)

                # Save the parsed JSON to a file for debugging
                with open("parsed_resume.json", "w") as f:
                    json.dump(analysis_data, indent=2, fp=f)

                # Add additional processing metrics
                analysis_data['technical_score'] = self._calculate_technical_score(
                    analysis_data)
                analysis_data['experience_years'] = self._calculate_total_experience(
                    analysis_data)
                analysis_data['education_score'] = self._calculate_education_score(
                    analysis_data)

                return analysis_data

            except json.JSONDecodeError as e:
                # If JSON parsing fails, return structured data with the raw text
                self.logger.error(
                    f"Failed to parse JSON from OpenAI response: {str(e)}")
                return {
                    "error": f"Failed to parse response: {str(e)}",
                    "raw_analysis": analysis_text,
                    "candidate_name": "Unknown",
                    "skills": [],
                    "experience": [],
                    "education": [],
                    "summary": "Error processing resume"
                }

        except Exception as e:
            self.logger.error(f"Error in analyze_resume: {str(e)}")
            return {
                "error": str(e),
                "candidate_name": "Error",
                "skills": [],
                "experience": [],
                "education": [],
                "summary": f"Error analyzing resume: {str(e)}"
            }

    def generate_job_description_match(self, resume_text: str, job_description: str) -> Dict[str, any]:
        """Generate a match analysis between resume and job description."""
        try:
            self.logger.info("Generating job description match")
            self.logger.info(
                f"Resume text length: {len(resume_text)} characters")
            self.logger.info(
                f"Job description length: {len(job_description)} characters")

            # Validate inputs
            if not resume_text or len(resume_text.strip()) < 100:
                self.logger.error("Resume text is too short or empty")
                return {
                    "error": "Resume text is too short or empty",
                    "overall_match_score": 0,
                    "skills_match": [],
                    "experience_relevance": 0,
                    "education_relevance": 0,
                    "strengths": [],
                    "weaknesses": [],
                    "recommendations": [],
                    "summary": "Cannot analyze match: resume text is insufficient"
                }

            if not job_description or len(job_description.strip()) < 50:
                self.logger.error("Job description is too short or empty")
                return {
                    "error": "Job description is too short or empty",
                    "overall_match_score": 0,
                    "skills_match": [],
                    "experience_relevance": 0,
                    "education_relevance": 0,
                    "strengths": [],
                    "weaknesses": [],
                    "recommendations": [],
                    "summary": "Cannot analyze match: job description is insufficient"
                }

            # Create prompt for OpenAI
            prompt = f"""
            Analyze how well the following resume matches the job description:
            
            RESUME:
            {resume_text}
            
            JOB DESCRIPTION:
            {job_description}
            
            Please provide a detailed analysis in JSON format with the following structure:
            {{
                "overall_match_score": 85,
                "skills_match": [
                    {{"skill": "Python", "match_percentage": 90, "level": "Advanced"}},
                    {{"skill": "Machine Learning", "match_percentage": 85, "level": "Advanced"}}
                ],
                "experience_relevance": 80,
                "education_relevance": 90,
                "strengths": ["Strong technical skills", "Relevant experience"],
                "weaknesses": ["Missing specific technology X"],
                "recommendations": ["Highlight project Y more prominently"],
                "summary": "Detailed analysis of the candidate's fit for the role, including career alignment and specific suggestions"
            }}
            
            For the summary field, provide a comprehensive 3-4 paragraph analysis that includes:
            1. An overall assessment of how well the candidate's background aligns with the job requirements
            2. Specific examples from the resume that demonstrate relevant experience or skills
            3. Areas where the candidate's experience differs from the job requirements
            4. Clear recommendations for how the candidate could better position themselves for this role
            5. Whether the candidate appears to be a career match for this position or might be better suited for a different role
            
            Be honest and accurate in your assessment. The overall_match_score should be a number between 0 and 100.
            Ensure your response is in valid JSON format only, with no additional text before or after the JSON.
            """

            # Call OpenAI API
            try:
                response = self.client.chat.completions.create(
                    model=self.model,
                    messages=[
                        {"role": "system", "content": "You are an expert HR analyst who evaluates resumes against job descriptions. Always respond with valid JSON only."},
                        {"role": "user", "content": prompt}
                    ],
                    temperature=0.3,
                    max_tokens=2000
                )

                # Extract and parse the response
                analysis_text = response.choices[0].message.content.strip()

                # Save the raw response to a file for debugging
                with open("job_match_response.txt", "w") as f:
                    f.write(analysis_text)

                # Try to extract JSON from the response
                try:
                    # Find JSON in the response if it's wrapped in markdown or other text
                    json_start = analysis_text.find('{')
                    json_end = analysis_text.rfind('}') + 1

                    if json_start >= 0 and json_end > json_start:
                        json_str = analysis_text[json_start:json_end]
                        self.logger.info(
                            f"Extracted JSON from response: {json_str[:100]}...")
                        match_data = json.loads(json_str)
                    else:
                        # Try to parse the entire response as JSON
                        self.logger.info(
                            "Attempting to parse entire response as JSON")
                        match_data = json.loads(analysis_text)

                    # Save the parsed JSON to a file for debugging
                    with open("job_match.json", "w") as f:
                        json.dump(match_data, indent=2, fp=f)

                    # Validate the required fields
                    if "overall_match_score" not in match_data:
                        self.logger.warning(
                            "Missing overall_match_score in response, setting to 0")
                        match_data["overall_match_score"] = 0

                    if "skills_match" not in match_data:
                        self.logger.warning(
                            "Missing skills_match in response, setting to empty list")
                        match_data["skills_match"] = []

                    if "experience_relevance" not in match_data:
                        self.logger.warning(
                            "Missing experience_relevance in response, setting to 0")
                        match_data["experience_relevance"] = 0

                    if "education_relevance" not in match_data:
                        self.logger.warning(
                            "Missing education_relevance in response, setting to 0")
                        match_data["education_relevance"] = 0

                    return match_data

                except json.JSONDecodeError as e:
                    # If JSON parsing fails, return structured data with the raw text
                    self.logger.error(
                        f"Failed to parse JSON from OpenAI response: {str(e)}")
                    self.logger.error(f"Raw response: {analysis_text}")

                    # Try to generate a fallback analysis
                    return self._generate_fallback_analysis(resume_text, job_description, analysis_text)

            except Exception as api_error:
                self.logger.error(f"OpenAI API error: {str(api_error)}")
                return {
                    "error": f"OpenAI API error: {str(api_error)}",
                    "overall_match_score": 0,
                    "skills_match": [],
                    "experience_relevance": 0,
                    "education_relevance": 0,
                    "strengths": [],
                    "weaknesses": [],
                    "recommendations": [],
                    "summary": "Error communicating with AI service"
                }

        except Exception as e:
            self.logger.error(
                f"Error in generate_job_description_match: {str(e)}")
            import traceback
            self.logger.error(traceback.format_exc())
            return {
                "error": str(e),
                "overall_match_score": 0,
                "skills_match": [],
                "experience_relevance": 0,
                "education_relevance": 0,
                "strengths": [],
                "weaknesses": [],
                "recommendations": [],
                "summary": "Error analyzing match"
            }

    def _generate_fallback_analysis(self, resume_text: str, job_description: str, raw_response: str) -> Dict[str, any]:
        """Generate a fallback analysis when JSON parsing fails."""
        self.logger.info("Generating fallback analysis")

        # Create a basic fallback analysis
        fallback = {
            "error": "Failed to parse AI response",
            "raw_analysis": raw_response[:500] + "..." if len(raw_response) > 500 else raw_response,
            "overall_match_score": 50,  # Default to 50% match
            "skills_match": [],
            "experience_relevance": 50,
            "education_relevance": 50,
            "strengths": ["Unable to determine strengths due to processing error"],
            "weaknesses": ["Unable to determine weaknesses due to processing error"],
            "recommendations": ["Try reprocessing the resume or check system logs for details"],
            "summary": """
The system encountered an error while analyzing this resume against the job requirements. 
A basic analysis has been generated based on keyword matching, but it may not reflect the full qualifications of the candidate.

Based on the limited analysis, the candidate appears to have some relevant skills for the position, but a more detailed review by a human recruiter is recommended. 
The system was unable to properly evaluate the candidate's experience and education relevance to the job requirements.

We recommend manually reviewing the resume to assess the candidate's qualifications, or try reprocessing the resume to get a more accurate analysis.
If this error persists, please contact technical support for assistance.
            """.strip()
        }

        # Try to extract some basic skills from the resume
        try:
            # Simple keyword extraction for skills
            common_skills = ["python", "java", "javascript", "html", "css", "react", "angular",
                             "node", "sql", "nosql", "mongodb", "aws", "azure", "docker",
                             "kubernetes", "ci/cd", "agile", "scrum", "project management",
                             "leadership", "communication", "teamwork", "problem solving",
                             "critical thinking", "time management", "customer service",
                             "sales", "marketing", "finance", "accounting", "hr", "human resources"]

            found_skills = []
            for skill in common_skills:
                if skill.lower() in resume_text.lower():
                    found_skills.append({
                        "skill": skill.capitalize(),
                        "match_percentage": 70,
                        "level": "Intermediate"
                    })

            if found_skills:
                # Limit to 5 skills
                fallback["skills_match"] = found_skills[:5]

                # Update summary with found skills
                skill_names = [skill["skill"] for skill in found_skills[:5]]
                if skill_names:
                    skills_text = ", ".join(skill_names)
                    fallback["summary"] += f"\n\nThe candidate appears to have skills in {skills_text}, which may be relevant to the position."

            # Try to determine if the resume is more aligned with a different role
            tech_keywords = ["python", "java", "javascript", "programming",
                             "developer", "software", "code", "algorithm", "data structure"]
            data_keywords = ["data", "analytics", "statistics", "machine learning",
                             "ai", "artificial intelligence", "model", "prediction"]
            marketing_keywords = ["marketing", "social media",
                                  "campaign", "brand", "content", "seo", "ppc", "advertising"]
            hr_keywords = ["hr", "human resources", "recruitment", "talent",
                           "onboarding", "employee", "benefits", "compensation"]
            finance_keywords = ["finance", "accounting", "budget",
                                "financial", "tax", "audit", "investment", "revenue"]

            tech_count = sum(
                1 for kw in tech_keywords if kw.lower() in resume_text.lower())
            data_count = sum(
                1 for kw in data_keywords if kw.lower() in resume_text.lower())
            marketing_count = sum(
                1 for kw in marketing_keywords if kw.lower() in resume_text.lower())
            hr_count = sum(1 for kw in hr_keywords if kw.lower()
                           in resume_text.lower())
            finance_count = sum(
                1 for kw in finance_keywords if kw.lower() in resume_text.lower())

            # Get the job role from the job description
            job_role = ""
            if "Job Title:" in job_description:
                job_role = job_description.split(
                    "Job Title:")[1].split("\n")[0].strip()

            # Determine if resume might be better suited for a different role
            counts = {
                "Software Development": tech_count,
                "Data Science": data_count,
                "Marketing": marketing_count,
                "Human Resources": hr_count,
                "Finance": finance_count
            }

            max_category = max(counts, key=counts.get)
            max_count = counts[max_category]

            if max_count > 5 and job_role and max_category.lower() not in job_role.lower():
                fallback["summary"] += f"\n\nBased on keyword analysis, this resume appears to be more aligned with {max_category} roles rather than the {job_role} position. Consider discussing with the candidate if they have specific interest or experience in {job_role} that may not be apparent from their resume."

        except Exception as skill_error:
            self.logger.error(
                f"Error extracting fallback skills: {str(skill_error)}")

        return fallback

    def _calculate_technical_score(self, analysis_data: Dict) -> float:
        """Calculate technical score based on skills."""
        try:
            if not analysis_data.get('skills'):
                return 0

            # Map skill levels to scores
            level_scores = {
                "beginner": 25,
                "intermediate": 50,
                "advanced": 75,
                "expert": 100
            }

            # Calculate average skill score
            total_score = 0
            count = 0

            for skill in analysis_data['skills']:
                level = skill.get('level', '').lower()
                # Default to intermediate if level not specified
                score = level_scores.get(level, 50)
                total_score += score
                count += 1

            return round(total_score / max(count, 1), 1)
        except Exception as e:
            self.logger.error(f"Error calculating technical score: {str(e)}")
            return 50  # Default to middle score

    def _calculate_total_experience(self, analysis_data: Dict) -> float:
        """Calculate total years of experience."""
        try:
            if not analysis_data.get('experience'):
                return 0

            total_years = 0

            for exp in analysis_data['experience']:
                duration_str = exp.get('duration', '0')
                try:
                    # Try to extract years from strings like "2 years" or "2.5 years"
                    years = float(
                        ''.join(c for c in duration_str if c.isdigit() or c == '.'))
                    total_years += years
                except ValueError:
                    continue

            return round(total_years, 1)
        except Exception as e:
            self.logger.error(f"Error calculating total experience: {str(e)}")
            return 0

    def _calculate_education_score(self, analysis_data: Dict) -> float:
        """Calculate education score."""
        try:
            if not analysis_data.get('education'):
                return 0

            # Map degree levels to scores
            degree_scores = {
                "high school": 25,
                "associate": 50,
                "bachelor": 75,
                "master": 90,
                "phd": 100,
                "doctorate": 100
            }

            max_score = 0

            for edu in analysis_data['education']:
                degree = edu.get('degree', '').lower()

                # Check for degree level keywords
                for level, score in degree_scores.items():
                    if level in degree:
                        max_score = max(max_score, score)
                        break

            return max_score
        except Exception as e:
            self.logger.error(f"Error calculating education score: {str(e)}")
            return 50  # Default to middle score


class ResumeProcessor:
    """Process resume files and extract text."""

    def __init__(self):
        self.logger = logging.getLogger(__name__)

    def extract_text_from_file(self, file_path: str, mime_type: str = None) -> str:
        """Extract text from various file formats."""
        try:
            self.logger.info(f"Extracting text from file: {file_path}")

            if not mime_type:
                # Determine MIME type if not provided
                import magic
                mime_type = magic.from_file(file_path, mime=True)

            self.logger.info(f"File MIME type: {mime_type}")

            # Extract text based on file type
            if mime_type == 'application/pdf':
                return self._extract_text_from_pdf(file_path)
            elif mime_type in ['application/vnd.openxmlformats-officedocument.wordprocessingml.document', 'application/msword']:
                return self._extract_text_from_docx(file_path)
            elif mime_type == 'text/plain':
                return self._extract_text_from_txt(file_path)
            else:
                error_msg = f"Unsupported file type: {mime_type}"
                self.logger.error(error_msg)
                return f"ERROR EXTRACTING TEXT: {error_msg}"

        except Exception as e:
            error_msg = f"Error extracting text: {str(e)}"
            self.logger.error(error_msg)
            import traceback
            self.logger.error(traceback.format_exc())
            return f"ERROR EXTRACTING TEXT: {error_msg}"

    def _extract_text_from_pdf(self, file_path: str) -> str:
        """Extract text from PDF files using multiple methods for reliability."""
        try:
            self.logger.info("Extracting text from PDF")

            # Try pdfplumber first (better for formatted PDFs)
            try:
                import pdfplumber
                text = ""
                with pdfplumber.open(file_path) as pdf:
                    for page in pdf.pages:
                        page_text = page.extract_text() or ""
                        text += page_text + "\n\n"

                if text.strip():
                    self.logger.info(
                        f"Successfully extracted text with pdfplumber: {len(text)} characters")
                    return text
            except Exception as e:
                self.logger.warning(
                    f"pdfplumber extraction failed: {str(e)}, trying PyPDF2")

            # Fallback to PyPDF2
            try:
                import PyPDF2
                text = ""
                with open(file_path, 'rb') as file:
                    reader = PyPDF2.PdfReader(file)
                    for page_num in range(len(reader.pages)):
                        page = reader.pages[page_num]
                        text += page.extract_text() + "\n\n"

                if text.strip():
                    self.logger.info(
                        f"Successfully extracted text with PyPDF2: {len(text)} characters")
                    return text
            except Exception as e:
                self.logger.warning(f"PyPDF2 extraction failed: {str(e)}")

            # If both methods failed, try a more aggressive approach
            try:
                import subprocess
                import tempfile

                # Use pdftotext (from poppler-utils) if available
                text = subprocess.check_output(
                    ['pdftotext', file_path, '-'], stderr=subprocess.DEVNULL).decode('utf-8', errors='ignore')

                if text.strip():
                    self.logger.info(
                        f"Successfully extracted text with pdftotext: {len(text)} characters")
                    return text
            except Exception as e:
                self.logger.warning(f"pdftotext extraction failed: {str(e)}")

            # If we got here, all methods failed
            error_msg = "All PDF extraction methods failed"
            self.logger.error(error_msg)
            return f"ERROR EXTRACTING TEXT: {error_msg}"

        except Exception as e:
            error_msg = f"Error in PDF extraction: {str(e)}"
            self.logger.error(error_msg)
            return f"ERROR EXTRACTING TEXT: {error_msg}"

    def _extract_text_from_docx(self, file_path: str) -> str:
        """Extract text from DOCX files."""
        try:
            self.logger.info("Extracting text from DOCX")

            import docx
            doc = docx.Document(file_path)

            # Extract text from paragraphs
            text = "\n\n".join(
                [para.text for para in doc.paragraphs if para.text.strip()])

            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(
                        [cell.text for cell in row.cells if cell.text.strip()])
                    if row_text:
                        text += "\n" + row_text

            self.logger.info(
                f"Successfully extracted text from DOCX: {len(text)} characters")
            return text

        except Exception as e:
            error_msg = f"Error in DOCX extraction: {str(e)}"
            self.logger.error(error_msg)
            return f"ERROR EXTRACTING TEXT: {error_msg}"

    def _extract_text_from_txt(self, file_path: str) -> str:
        """Extract text from plain text files."""
        try:
            self.logger.info("Extracting text from TXT")

            with open(file_path, 'r', encoding='utf-8', errors='ignore') as file:
                text = file.read()

            self.logger.info(
                f"Successfully extracted text from TXT: {len(text)} characters")
            return text

        except Exception as e:
            error_msg = f"Error in TXT extraction: {str(e)}"
            self.logger.error(error_msg)
            return f"ERROR EXTRACTING TEXT: {error_msg}"
